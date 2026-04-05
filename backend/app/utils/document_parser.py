"""Document parsing: PDF, DOCX, TXT, and image OCR extraction."""

import base64
import io
import os
import re
from typing import Dict, Tuple

import PyPDF2
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageFilter, ImageEnhance
import requests


class DocumentParser:
    @staticmethod
    def extract_from_base64(file_base64: str, file_type: str) -> str:
        try:
            file_bytes = base64.b64decode(file_base64)
        except Exception as e:
            raise ValueError(f"Invalid base64 encoding: {e}")

        file_type = file_type.lower().strip()

        if file_type == "pdf":
            return DocumentParser._from_pdf(file_bytes)
        elif file_type == "docx":
            return DocumentParser._from_docx(file_bytes)
        elif file_type in ("txt", "text"):
            return DocumentParser._from_text(file_bytes)
        elif file_type in ("image", "jpg", "jpeg", "png", "gif", "bmp"):
            return DocumentParser._from_image(file_bytes)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    @staticmethod
    def _from_pdf(file_bytes: bytes) -> str:
        pdf_file = io.BytesIO(file_bytes)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page_num, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text() or ""
            text += f"\n{'='*60}\n"
            text += f"  PAGE {page_num + 1} OF {len(pdf_reader.pages)}\n"
            text += f"{'='*60}\n\n"
            text += DocumentParser._detect_pdf_structure(page_text)
        return text.strip()

    @staticmethod
    def _detect_pdf_structure(page_text: str) -> str:
        lines = page_text.split('\n')
        result = []
        for line in lines:
            stripped = line.strip()
            if not stripped:
                result.append("")
                continue
            if '|' in stripped and stripped.count('|') >= 2:
                result.append(f"[TABLE] {stripped} [/TABLE]")
            elif stripped.isupper() and len(stripped) < 80 and len(stripped.split()) <= 10:
                result.append(f"\n## {stripped}\n")
            elif re.match(r'^[\-\*\u2022]\s+', stripped):
                result.append(f"  - {stripped.lstrip('-*• ').strip()}")
            else:
                result.append(stripped)
        return '\n'.join(result)

    @staticmethod
    def _from_docx(file_bytes: bytes) -> str:
        docx_file = io.BytesIO(file_bytes)
        doc = Document(docx_file)
        text = ""
        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                text += "\n"
                continue
            style_name = paragraph.style.name.lower() if paragraph.style else ""
            if 'heading 1' in style_name:
                text += f"\n{'='*60}\n  {paragraph.text}\n{'='*60}\n\n"
            elif 'heading 2' in style_name:
                text += f"\n## {paragraph.text}\n\n"
            elif 'heading 3' in style_name:
                text += f"\n### {paragraph.text}\n\n"
            elif 'list' in style_name:
                text += f"  - {paragraph.text}\n"
            elif paragraph.runs and any(run.bold for run in paragraph.runs if run.text.strip()):
                text += f"\n**{paragraph.text}**\n\n"
            else:
                text += paragraph.text + "\n"
        if doc.tables:
            text += f"\n{'─'*60}\n  TABLES\n{'─'*60}\n\n"
            for idx, table in enumerate(doc.tables):
                text += f"[Table {idx + 1}]\n"
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        text += " | ".join(cells) + "\n"
                text += "\n"
        return text.strip()

    @staticmethod
    def _from_text(file_bytes: bytes) -> str:
        return file_bytes.decode('utf-8', errors='ignore').strip()

    @staticmethod
    def _from_image(file_bytes: bytes) -> str:
        """Use OCR.space free API for image OCR."""
        api_key = os.getenv('OCR_SPACE_API_KEY', 'helloworld')
        url = 'https://api.ocr.space/parse/image'
        payload = {
            'apikey': api_key,
            'language': 'eng',
            'isOverlayRequired': 'false',
            'scale': 'true',
            'OCREngine': '2',
        }
        files = {'file': ('image.png', io.BytesIO(file_bytes), 'image/png')}
        resp = requests.post(url, data=payload, files=files, timeout=60)
        resp.raise_for_status()
        data = resp.json()
        if data.get('IsErroredOnProcessing'):
            raise RuntimeError(f"OCR.space error: {data.get('ErrorMessage', 'Unknown error')}")
        parsed_results = data.get('ParsedResults', [])
        if not parsed_results:
            return ""
        return parsed_results[0].get('ParsedText', '').strip()

    @staticmethod
    def _preprocess_image(image: Image.Image) -> Image.Image:
        if image.mode != 'L':
            image = image.convert('L')
        image = image.point(lambda x: 0 if x < 128 else 255, '1')
        image = image.filter(ImageFilter.MedianFilter(size=3))
        scale_factor = max(1, 2000 / max(image.size))
        if scale_factor > 1:
            new_size = (int(image.width * scale_factor), int(image.height * scale_factor))
            image = image.resize(new_size, Image.LANCZOS)
        return image

    @staticmethod
    def get_metadata(file_bytes: bytes, file_type: str) -> Dict:
        metadata = {"file_type": file_type, "size_bytes": len(file_bytes)}
        if file_type == "pdf":
            try:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
                metadata["page_count"] = len(pdf_reader.pages)
            except Exception:
                pass
        elif file_type == "docx":
            try:
                doc = Document(io.BytesIO(file_bytes))
                metadata["paragraph_count"] = len(doc.paragraphs)
                metadata["table_count"] = len(doc.tables)
            except Exception:
                pass
        elif file_type in ("image", "jpg", "jpeg", "png", "gif", "bmp"):
            try:
                img = Image.open(io.BytesIO(file_bytes))
                metadata["width"] = img.width
                metadata["height"] = img.height
            except Exception:
                pass
        return metadata


def extract_text(file_base64: str, file_type: str) -> Tuple[str, bool]:
    try:
        text = DocumentParser.extract_from_base64(file_base64, file_type)
        return text, True
    except Exception as e:
        return str(e), False
